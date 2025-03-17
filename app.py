import os
import re
import PyPDF2
import docx
import google.generativeai as genai
import json
from io import BytesIO
from flask import Flask, request, render_template, jsonify, send_file, redirect, url_for, flash
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF for PDF processing
from docx import Document

app = Flask(__name__)
app.secret_key = 'resume_formatter_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'doc'}

# Create upload folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize Gemini API with hardcoded key
gemini_api_key = "your api key"  # Your Gemini API key
gemini_model = None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyMuPDF for better results"""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def identify_sections(text):
    """Use Gemini API to identify resume sections"""
    global gemini_model
    
    if not gemini_model:
        return {"error": "Gemini API not configured"}
    
    try:
        prompt = f"""
        Please analyze this resume text and identify all the different sections.
        Return the result as a JSON object with this format:
        {{
            "sections": [
                {{"name": "SECTION_NAME", "content": "SECTION_CONTENT"}}
            ]
        }}
        
        Here's the resume text:
        {text}
        """
        
        response = gemini_model.generate_content(prompt)
        content = response.text
        
        # Try to extract JSON from the response
        json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)
        
        parsed = json.loads(content)
        return parsed
    except Exception as e:
        return {"error": f"Error identifying sections: {str(e)}", "sections": [{"name": "Full Resume", "content": text}]}

def format_resume(text, sections_to_include=None, format_instructions=None):
    """Use Gemini API to format the resume"""
    global gemini_model
    
    if not gemini_model:
        return {"error": "Gemini API not configured"}
    
    try:
        if sections_to_include:
            section_text = "Include only these sections: " + ", ".join(sections_to_include)
        else:
            section_text = "Include all sections"
            
        instructions = format_instructions or "Format the resume professionally with clear section headings"
        
        prompt = f"""
        Please format this resume according to the following instructions:
        {instructions}
        {section_text}
        
        Here's the resume content:
        {text}
        
        Return the formatted resume as plain text with proper spacing and formatting.
        """
        
        response = gemini_model.generate_content(prompt)
        return {"formatted_text": response.text}
    except Exception as e:
        return {"error": f"Error formatting resume: {str(e)}"}

def create_docx(text):
    """Create a Word document with formatted text"""
    doc = Document()
    for line in text.split('\n'):
        if line.strip():
            # Check if line is a heading (all caps or ends with :)
            if line.strip().isupper() or line.strip().endswith(':'):
                doc.add_heading(line.strip(), level=2)
            else:
                doc.add_paragraph(line.strip())
        else:
            # Add empty paragraph for line breaks
            doc.add_paragraph('')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def create_pdf(text):
    """Create a PDF document from formatted text using PyMuPDF"""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/')
def index():
    global gemini_api_key
    api_configured = gemini_api_key is not None
    return render_template('index.html', api_configured=api_configured)

@app.route('/configure_api', methods=['POST'])
def configure_api():
    global gemini_api_key, gemini_model
    
    api_key = request.form.get('api_key')
    if not api_key:
        flash('API key is required')
        return redirect(url_for('index'))
    
    try:
        # Configure Gemini API
        genai.configure(api_key=api_key)
        # Test the API
        generation_config = {
            "temperature": 0.7,
            "top_p": 0.95,
            "top_k": 40,
        }
        gemini_model = genai.GenerativeModel(
            model_name="gemini-1.5-pro",
            generation_config=generation_config
        )
        # Simple test to verify API works
        test_response = gemini_model.generate_content("Hello")
        if test_response:
            gemini_api_key = api_key
            flash('Gemini API configured successfully')
        else:
            flash('Failed to test Gemini API')
    except Exception as e:
        flash(f'Error configuring Gemini API: {str(e)}')
    
    return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('index'))
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('index'))
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith(('.docx', '.doc')):
            text = extract_text_from_docx(file_path)
        else:
            text = "Unsupported file format"
        
        # Identify sections
        sections_result = identify_sections(text)
        
        return render_template(
            'sections.html', 
            filename=filename, 
            text=text, 
            sections=sections_result.get('sections', []),
            error=sections_result.get('error')
        )
    
    flash('Invalid file type. Please upload PDF or Word documents.')
    return redirect(url_for('index'))

@app.route('/format', methods=['POST'])
def format_resume_route():
    text = request.form.get('text', '')
    output_format = request.form.get('output_format', 'pdf')
    format_instructions = request.form.get('format_instructions', '')
    
    sections_to_include = request.form.getlist('sections')
    
    result = format_resume(text, sections_to_include, format_instructions)
    
    if 'error' in result:
        flash(result['error'])
        return redirect(url_for('index'))
    
    formatted_text = result['formatted_text']
    
    # Create output file
    if output_format.lower() == 'docx':
        output = create_docx(formatted_text)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='formatted_resume.docx'
        )
    else:  # Default to PDF
        output = create_pdf(formatted_text)
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='formatted_resume.pdf'
        )

@app.route('/extract_section', methods=['POST'])
def extract_section():
    section_name = request.form.get('section_name')
    section_content = request.form.get('section_content')
    
    return jsonify({
        'section_name': section_name,
        'section_content': section_content
    })

if __name__ == '__main__':
    # Configure Gemini API on startup
    try:
        genai.configure(api_key=gemini_api_key)
        generation_config = {
            "temperature": 0.7,
            "top_p": 0.95,
            "top_k": 40,
        }
        gemini_model = genai.GenerativeModel(
            model_name="gemini-1.5-pro",
            generation_config=generation_config
        )
        print("Gemini API configured successfully on startup")
    except Exception as e:
        print(f"Error configuring Gemini API: {str(e)}")
    
    app.run(debug=True)